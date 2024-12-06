import streamlit as st
from dotenv import load_dotenv
from TextAnalytics.utils import *
import spacy
from spacy import displacy  # visualization
from TextAnalytics.helper import *
from io import BytesIO
from docx import Document

# HTML Wrapper for results display
HTML_WRAPPER = """<div style="overflow-x: auto; border: 1px solid #e6e9ef; border-radius: 0.25rem; padding: 1rem">{}</div>"""

# Input data handler
def get_input_data():
    uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")
    url = st.text_input("Enter a URL")
    text = st.text_area("Enter your text", height=200)

    if uploaded_file:
        extracted_text = extract_text_from_pdf(uploaded_file)
        input_type = "PDF"
    elif url:
        extracted_text = extract_text_from_url(url)
        input_type = "URL"
    elif text:
        extracted_text = text
        input_type = "Text"
    else:
        extracted_text = None
        input_type = None

    return extracted_text, input_type

# Set up Streamlit page
st.set_page_config(layout="wide")
st.title("Text Analytics")

# Get the input data from the user
input_text, input_type = get_input_data()

# Check if the Analyze Text button was clicked
if st.button("Analyze Text:"):
    if input_text:
        # Display the extracted text
        st.text_area(f"Extracted Text from {input_type}", input_text, height=200)

        # Key Findings based on Text (Summarization)
        st.markdown("**Key Findings based on your Text**")
        key_findings = extract_key_findings(input_text)
        if key_findings:
            st.success(key_findings)
        else:
            st.warning("Could not extract key findings.")

        # Word Cloud Generation
        st.markdown("**Output Image (Word Cloud)**")
        wordcloud_image_path = generate_wordcloud(input_text)
        if wordcloud_image_path:
            st.image(wordcloud_image_path)
        else:
            st.warning("Could not generate word cloud.")

        # Sentiment Analysis
        st.markdown("**Sentiment Analysis**")
        sentiment = sentiment_analysis(input_text)
        if sentiment:
            st.success(sentiment)
        else:
            st.warning("Could not perform sentiment analysis.")

        # Most Positive Words
        st.markdown("**Most Positive Words**")
        positive_words = most_positive_words(input_text)
        if positive_words:
            st.success(positive_words)
        else:
            st.warning("Could not extract most positive words.")

        # Summarization
        st.markdown("**Summary**")
        summary = summarization(input_text)
        if summary:
            st.success(summary)
        else:
            st.warning("Could not generate a summary.")

        # Optional: Named Entity Recognition (NER) visualization
        st.markdown("**Named Entity Recognition (NER)**")
        nlp = spacy.load("en_core_web_sm")  # Load SpaCy model
        doc = nlp(input_text)
        html = displacy.render(doc, style="ent", page=True)
        st.markdown(HTML_WRAPPER.format(html), unsafe_allow_html=True)

        # Combine all results for download
        all_results = f"**Key Findings**\n{key_findings}\n\n**Sentiment Analysis**\n{sentiment}\n\n**Summary**\n{summary}\n\n**Word Cloud Image Path**\n{wordcloud_image_path}\n\n**Most Positive Words**\n{positive_words}\n\n**NER Visualization**\n{html}"

        # Convert to TXT and DOCX
        def convert_to_txt(results):
            return results.encode("utf-8")

        def convert_to_docx(results):
            doc = Document()
            doc.add_heading("Text Analytics Results", 0)

            # Add each section
            doc.add_heading("Key Findings", level=1)
            doc.add_paragraph(key_findings)

            doc.add_heading("Sentiment Analysis", level=1)
            doc.add_paragraph(sentiment)

            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary)

            doc.add_heading("Word Cloud Image Path", level=1)
            doc.add_paragraph(wordcloud_image_path)

            doc.add_heading("Most Positive Words", level=1)
            doc.add_paragraph(positive_words)

            doc.add_heading("NER Visualization", level=1)
            doc.add_paragraph(html)

            # Embed wordcloud image in Word document
            doc.add_heading("Word Cloud Image", level=1)
            doc.add_paragraph("See below for the generated word cloud:")
            doc.add_picture(wordcloud_image_path)

            # Save doc to BytesIO buffer
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            return doc_buffer

        # Download buttons for result files
        st.download_button(
            label="Download as TXT",
            data=convert_to_txt(all_results),
            file_name="summarized_text.txt",
            mime="text/plain",
        )
        st.download_button(
            label="Download as DOCX",
            data=convert_to_docx(all_results),
            file_name="summarized_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.warning("Please upload or enter some text to analyze.")
